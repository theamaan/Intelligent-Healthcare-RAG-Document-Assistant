import re
from pathlib import Path

import fitz
from docx import Document as DocxDocument
from langchain_core.documents import Document


ALLOWED_EXTENSIONS = {".pdf", ".docx"}


def clean_text(text: str) -> str:
    text = text.replace("\x00", " ")
    text = re.sub(r"\s+", " ", text)
    return text.strip()


def load_pdf(file_path: Path, doc_id: str) -> list[Document]:
    documents: list[Document] = []
    with fitz.open(file_path) as pdf:
        for page_index, page in enumerate(pdf):
            text = clean_text(page.get_text())
            if not text:
                continue
            documents.append(
                Document(
                    page_content=text,
                    metadata={
                        "doc_id": doc_id,
                        "source": file_path.name,
                        "page_number": page_index + 1,
                        "section": f"page_{page_index + 1}",
                    },
                )
            )
    return documents


def load_docx(file_path: Path, doc_id: str) -> list[Document]:
    documents: list[Document] = []
    doc = DocxDocument(str(file_path))
    for index, paragraph in enumerate(doc.paragraphs):
        text = clean_text(paragraph.text)
        if not text:
            continue
        documents.append(
            Document(
                page_content=text,
                metadata={
                    "doc_id": doc_id,
                    "source": file_path.name,
                    "page_number": None,
                    "section": f"paragraph_{index + 1}",
                },
            )
        )
    return documents


def load_document(file_path: Path, doc_id: str) -> list[Document]:
    extension = file_path.suffix.lower()
    if extension not in ALLOWED_EXTENSIONS:
        raise ValueError("Unsupported file type. Only PDF and DOCX are allowed.")
    if extension == ".pdf":
        return load_pdf(file_path, doc_id)
    return load_docx(file_path, doc_id)
